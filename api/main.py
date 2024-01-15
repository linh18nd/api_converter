from io import BytesIO
import logging
import os
import secrets
import subprocess
import uuid
from datetime import datetime, timedelta
from pathlib import Path
from threading import BoundedSemaphore
from typing import Optional, Dict, Set, Union
from uuid import UUID
from fastapi.middleware.cors import CORSMiddleware
from fastapi.params import Form
from fastapi.responses import FileResponse, StreamingResponse
from docx import Document as DocxDocument
from api.database import connect, disconnect, Base, execute
from api.database import connect, disconnect, Base, execute
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from fastapi import (
    FastAPI,
    File,
    UploadFile,
    HTTPException,
    BackgroundTasks,
    Depends,
    Security,
    Response,
    Query,
)
from fastapi.openapi.models import APIKey
from fastapi.responses import FileResponse, JSONResponse
from fastapi.security import APIKeyHeader
from starlette.status import HTTP_403_FORBIDDEN
from api import database
from api.models import Document, Lang

from fastapi import Query

from typing import List


from api.settings import config
from api.database import DBDocument
from api.tools import save_upload_file

logger = logging.getLogger("gunicorn.error")

app = FastAPI(
    title="api_converter",
    description="Basic API for OCR PDF to TXT",
    version="0.0.3",
    redoc_url=None,
)
Schedule = AsyncIOScheduler({"apscheduler.timezone": "UTC"})
Schedule.start()

pool_ocr = BoundedSemaphore(value=config.max_ocr_process)

document: Document
workdir = config.workdir

script_directory = Path(os.path.dirname(os.path.abspath(__file__))).resolve()
expiration_delta = timedelta(hours=config.document_expire_hour)


def delete_data(doc: DBDocument):
    output_path = Path(doc.output)
    output_json_path = Path(doc.output_json)
    output_txt_path = Path(doc.output_txt)
    input_path = Path(doc.input)
    if output_path.exists():
        output_path.unlink()
    if output_json_path.exists():
        output_json_path.unlink()
    if output_txt_path.exists():
        output_txt_path.unlink()
    if input_path.exists():
        input_path.unlink()


if not workdir.exists():
    workdir.mkdir()


async def do_ocr(_doc: Document):
    pool_ocr.acquire()
    _doc.ocr(config.enable_wsl_compat)
    pool_ocr.release()


def get_db():
    db = execute(DBDocument.__table__.metadata.tables['documents'].select())
    return db


# Định cấu hình CORS
origins = [
    "http://localhost",
    "http://localhost:8080",
    "http://localhost:8000",
    "https://example.com",
    "https://staging.example.com",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"], 
    allow_headers=["*"],
)


@app.on_event("startup")
async def startup_db_client():
    await database.connect()


@app.on_event("shutdown")
async def shutdown_db_client():
    await database.disconnect()

api_key_header = APIKeyHeader(name="X-API-KEY")


async def check_api_key(x_api_key: str = Security(api_key_header)):
    if secrets.compare_digest(x_api_key, config.api_key_secret):
        return api_key_header
    else:
        raise HTTPException(
            status_code=HTTP_403_FORBIDDEN, detail="Could not validate credentials"
        )


@app.get("/status", include_in_schema=False)
def status():
    ocrmypdf = subprocess.check_output(
        f"{config.base_command_ocr} --version", shell=True
    )
    return {"status": "ok", "version_ocr": ocrmypdf.strip()}


@app.get("/ocr/{pid}/pdf")
async def get_doc_pdf(pid: UUID, api_key: APIKey = Depends(check_api_key)):
    query = DBDocument.__table__.select().where(
        DBDocument.__table__.c.pid == str(pid))
    doc = await database.fetch_one(query)

    if doc:
        path = Path(doc.output)

        if path.resolve().exists():
            return FileResponse(
                str(doc.output),
                headers={"Content-Type": "application/pdf"},
                filename=f"{pid}.pdf",
            )

    raise HTTPException(status_code=404, detail="Document not found")


@app.get("/ocr/{pid}/txt")
async def get_doc_txt(pid: UUID, api_key: APIKey = Depends(check_api_key)):
    query = DBDocument.__table__.select().where(
        DBDocument.__table__.c.pid == str(pid))
    doc = await database.fetch_one(query)

    if doc:
        path = Path(doc.output_txt)

        if path.resolve().exists():
            return FileResponse(
                str(doc.output_txt),
                headers={"Content-Type": "text/plain"},
                filename=f"{pid}.txt",
            )

    raise HTTPException(status_code=404, detail="Document not found")


@app.get("/ocr/{pid}/docx")
async def get_doc_docx(pid: UUID, api_key: APIKey = Depends(check_api_key)):
    query = DBDocument.__table__.select().where(
        DBDocument.__table__.c.pid == str(pid))
    doc = await database.fetch_one(query)

    if doc:
        path = Path(doc.output_txt)

        if path.resolve().exists():
            doc = DocxDocument()
            with open(str(path.resolve()), 'r', encoding='utf-8') as txt_file:
                for line in txt_file:
                    doc.add_paragraph(line.strip())

            docx_content = BytesIO()
            doc.save(docx_content)

            docx_content.seek(0)

            return StreamingResponse(
                content=docx_content,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f"attachment; filename={pid}.docx"},
            )

    raise HTTPException(status_code=404)


# @app.get("/search", response_model=list)
# async def search_files(queries: List[str] = Query(..., title="Search Queries"), api_key: APIKey = Depends(check_api_key)):
#     docs = await database.fetch_all(DBDocument.__table__.select())
#     matching_pdfs = []

#     for doc in docs:
#         path = Path(doc.output_txt)
#         if path.exists():
#             with open(str(doc.output_txt), 'r', encoding='utf-8') as txt_file:
#                 txt_content = txt_file.read()

#                 if all(query.lower() in txt_content.lower() for query in queries):
#                     matching_pdfs.append({"pid": str(doc.pid), "file_name": f"{doc.file_name}.pdf"})

#     return matching_pdfs


def parse_search_query(query: str) -> List[Union[List[str], str]]:
    result = []
    groups = query.split("||")
    for group in groups:
        and_conditions = group.split("&&")
        if len(and_conditions) > 1:
            result.append(and_conditions)
        else:
            result.append([group])
    return result


def evaluate_condition(condition: List[str], txt_content: str) -> bool:
    return all(query.strip().lower() in txt_content.lower() for query in condition)


@app.get("/search", response_model=list)
async def search_files(
    search_query: str = Query(..., title="Search Query"),
    api_key: APIKey = Depends(check_api_key)
):
    docs = await database.fetch_all(DBDocument.__table__.select())
    matching_pdfs = []

    parsed_query = parse_search_query(search_query)

    for doc in docs:
        path = Path(doc.output_txt)
        if path.exists():
            with open(str(doc.output_txt), 'r', encoding='utf-8') as txt_file:
                txt_content = txt_file.read()

                # Evaluate each group in the parsed query
                if any(evaluate_condition(group, txt_content) for group in parsed_query):
                    matching_pdfs.append(
                        {"pid": str(doc.pid), "file_name": f"{doc.file_name}.pdf"})

    return matching_pdfs


@app.delete("/ocr/{pid}")
async def delete_doc(pid: UUID, api_key: APIKey = Depends(check_api_key)):
    query = DBDocument.__table__.select().where(
        DBDocument.__table__.c.pid == str(pid))
    doc = await database.fetch_one(query)
    if doc:
        delete_data(doc)
        await database.execute(DBDocument.__table__.delete().where(DBDocument.__table__.c.pid == str(pid)))
        return Response(status_code=204, headers={"X-Status": "Deleted"})
    raise HTTPException(status_code=404, detail="Document not found")


@app.get("/documents", response_model=list)
async def get_documents(api_key: APIKey = Depends(check_api_key)):
    documents = await database.fetch_all(DBDocument.__table__.select())
    result = [{"pid": doc.pid, "file_name": doc.file_name, "lang": doc.lang}
              for doc in documents]
    return result


@app.post(
    "/ocr", response_model=Document, status_code=200,
)
async def ocr(
    background_tasks: BackgroundTasks,
    lang: Optional[Set[str]] = Query([Lang.eng]),
    file: UploadFile = File(...),
    api_key: APIKey = Depends(check_api_key),
    file_name: Optional[str] = Query(None),

):
    pid = uuid.uuid4()
    now = datetime.now()
    expire = now + expiration_delta
    filename = f"{pid}"

    input_file = workdir / Path(f"i_{filename}.pdf")
    save_upload_file(file, input_file)
    output_file = workdir / Path(f"o_{filename}.pdf")
    output_file_json = workdir / Path(f"o_{filename}.json")
    output_file_txt = workdir / Path(f"o_{filename}.txt")

    document = Document(
        pid=pid,
        lang=lang,
        status="received",
        input=input_file,
        output=output_file,
        output_json=output_file_json,
        output_txt=output_file_txt,
        created=now,
        expire=expire,
        file_name=file_name or file.file,
    )

    await do_ocr(document)

    document.save_state()
    with open(str(output_file_txt), 'r', encoding='utf-8') as txt_file:
        txt_content = txt_file.read()
        await database.execute(
            DBDocument.__table__.insert().values(
                pid=str(pid),
                lang=",".join(lang),
                status="received",
                input=str(input_file.resolve()),
                output=str(output_file.resolve()),
                output_json=str(output_file_json.resolve()),
                output_txt=str(output_file_txt.resolve()),
                created=now,
                expire=expire,
                file_name=file_name or file.file,
                text=txt_content,
            )
        )

    return document
